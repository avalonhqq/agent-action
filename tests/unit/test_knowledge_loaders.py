from io import BytesIO

import pymupdf
from docx import Document

from bili_support.knowledge.loaders import create_default_loader_registry
from bili_support.knowledge.table_normalization import normalize_table
from bili_support.knowledge.types import SourceBlockType


def test_markdown_loader_preserves_heading_path_and_table_semantics() -> None:
    content = """# 大会员

## 自动续费

关闭后下月不再扣费。

| 套餐 | 价格 |
|---|---|
| 月卡 | 25元 |
""".encode()

    loaded = create_default_loader_registry().load(
        content=content,
        filename="membership.md",
        media_type="text/markdown",
    )

    table = next(
        block for block in loaded.blocks if block.block_type is SourceBlockType.TABLE
    )
    assert table.heading_path == ("大会员", "自动续费")
    assert "套餐=月卡" in table.content
    assert "价格=25元" in table.content


def test_docx_loader_keeps_document_order_and_table_headers() -> None:
    document = Document()
    document.add_heading("订单规则", level=1)
    document.add_paragraph("订单支付后可以查询物流。")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "状态"
    table.rows[0].cells[1].text = "说明"
    table.rows[1].cells[0].text = "已支付"
    table.rows[1].cells[1].text = "等待发货"
    buffer = BytesIO()
    document.save(buffer)

    loaded = create_default_loader_registry().load(
        content=buffer.getvalue(),
        filename="orders.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert [block.block_type for block in loaded.blocks] == [
        SourceBlockType.HEADING,
        SourceBlockType.PARAGRAPH,
        SourceBlockType.TABLE,
    ]
    assert loaded.blocks[2].heading_path == ("订单规则",)
    assert "状态=已支付" in loaded.blocks[2].content


def test_pdf_loader_records_page_number() -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "BiliSupport PDF knowledge")
    content = document.tobytes()
    document.close()

    loaded = create_default_loader_registry().load(
        content=content,
        filename="support.pdf",
        media_type="application/pdf",
    )

    assert loaded.metadata["page_count"] == 1
    assert loaded.blocks[0].page_number == 1
    assert "BiliSupport PDF knowledge" in loaded.blocks[0].content


def test_text_loader_supports_gb18030_and_paragraph_boundaries() -> None:
    loaded = create_default_loader_registry().load(
        content="第一段\n\n第二段".encode("gb18030"),
        filename="support.txt",
        media_type="text/plain",
    )

    assert [block.content for block in loaded.blocks] == ["第一段", "第二段"]


def test_single_cell_word_callout_is_not_repeated_as_its_own_header() -> None:
    normalized = normalize_table([["测试提示\n本文内容只用于解析测试。"]])

    assert normalized == "第1行：测试提示=本文内容只用于解析测试。"
    assert normalized.count("测试提示") == 1
