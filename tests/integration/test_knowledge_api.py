import sqlite3
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from bili_support.core.config import Settings
from bili_support.main import create_app


def _settings(tmp_path: Path) -> Settings:
    return Settings(
        _env_file=None,
        database_url=f"sqlite+aiosqlite:///{(tmp_path / 'knowledge.db').as_posix()}",
        database_auto_create=True,
        knowledge_storage_dir=str(tmp_path / "files"),
        api_token="knowledge-test-token",
        ui_enabled=False,
    )


def _headers(user_id: str = "knowledge-admin") -> dict[str, str]:
    return {
        "Authorization": "Bearer knowledge-test-token",
        "X-User-ID": user_id,
        "X-User-Name": user_id,
    }


def _upload(
    client: TestClient,
    content: bytes,
    *,
    filename: str = "membership.md",
) -> object:
    return client.post(
        "/api/v1/knowledge/documents",
        headers=_headers(),
        data={
            "title": "大会员规则",
            "business_domain": "membership",
            "access_scope": "public,support",
        },
        files={"file": (filename, content, "text/markdown")},
    )


def test_upload_is_idempotent_and_changed_content_creates_version(
    tmp_path: Path,
) -> None:
    with TestClient(create_app(_settings(tmp_path))) as client:
        first = _upload(client, b"# Membership\n\nVersion one.")
        duplicate = _upload(client, b"# Membership\n\nVersion one.")
        changed = _upload(client, b"# Membership\n\nVersion two.")
        document_id = first.json()["data"]["document"]["id"]
        versions = client.get(
            f"/api/v1/knowledge/documents/{document_id}/versions",
            headers=_headers(),
        )

    assert first.status_code == 201
    assert first.json()["data"]["job_status"] == "succeeded"
    assert first.json()["data"]["block_count"] == 2
    assert first.json()["data"]["chunk_count"] == 2
    assert duplicate.json()["data"]["deduplicated"] is True
    assert duplicate.json()["data"]["version"]["id"] == first.json()["data"]["version"]["id"]
    assert changed.json()["data"]["version"]["version_number"] == 2
    assert len(versions.json()["data"]) == 2


def test_failed_document_can_retry_and_remains_auditable(tmp_path: Path) -> None:
    with TestClient(create_app(_settings(tmp_path))) as client:
        failed = _upload(client, b"not-a-pdf", filename="broken.pdf")
        job_id = failed.json()["data"]["job_id"]
        retried = client.post(
            f"/api/v1/knowledge/jobs/{job_id}/retry",
            headers=_headers(),
        )

    assert failed.status_code == 201
    assert failed.json()["data"]["job_status"] == "failed"
    assert failed.json()["data"]["error_code"] == "DOCUMENT_SIGNATURE_MISMATCH"
    assert retried.status_code == 200
    assert retried.json()["data"]["job_status"] == "failed"
    assert retried.json()["data"]["attempt_count"] == 2


def test_document_is_isolated_by_owner_and_soft_deleted(tmp_path: Path) -> None:
    with TestClient(create_app(_settings(tmp_path))) as client:
        uploaded = _upload(client, b"# Account\n\nRecovery.")
        document_id = uploaded.json()["data"]["document"]["id"]
        hidden = client.get(
            f"/api/v1/knowledge/documents/{document_id}/versions",
            headers=_headers("other-admin"),
        )
        deleted = client.delete(
            f"/api/v1/knowledge/documents/{document_id}",
            headers=_headers(),
        )
        documents = client.get(
            "/api/v1/knowledge/documents",
            headers=_headers(),
        )

    assert hidden.status_code == 404
    assert deleted.status_code == 204
    assert documents.json()["data"] == []


def test_docx_faq_upload_persists_keywords_and_exposes_chunks(tmp_path: Path) -> None:
    document = Document()
    document.add_heading("客服FAQ", level=1)
    document.add_paragraph("Q：大会员开通后多久生效？")
    document.add_paragraph("A：支付成功后立即生效。")
    document.add_paragraph("关键词：生效时间、未到账、支付成功")
    buffer = BytesIO()
    document.save(buffer)

    with TestClient(create_app(_settings(tmp_path))) as client:
        uploaded = client.post(
            "/api/v1/knowledge/documents",
            headers=_headers(),
            data={
                "title": "大会员FAQ",
                "business_domain": "membership",
                "knowledge_type": "mixed",
                "access_scope": "public,support",
            },
            files={
                "file": (
                    "membership-faq.docx",
                    buffer.getvalue(),
                    (
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                )
            },
        )
        version_id = uploaded.json()["data"]["version"]["id"]
        chunks = client.get(
            f"/api/v1/knowledge/versions/{version_id}/chunks",
            headers=_headers(),
            params={"kind": "child"},
        )

    assert uploaded.status_code == 201
    assert uploaded.json()["data"]["chunk_count"] == 2
    assert chunks.status_code == 200
    assert len(chunks.json()["data"]) == 1
    child = chunks.json()["data"][0]
    assert child["metadata_json"]["strategy"] == "faq"
    assert child["metadata_json"]["keywords"] == [
        "生效时间",
        "未到账",
        "支付成功",
    ]


def test_duplicate_upload_backfills_chunks_created_before_5b(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    content = b"# Membership\n\nExisting source blocks without chunks."
    with TestClient(create_app(settings)) as client:
        _upload(client, content)

    database_path = tmp_path / "knowledge.db"
    with sqlite3.connect(database_path) as connection:
        connection.execute("DELETE FROM knowledge_chunks")
        connection.commit()

    with TestClient(create_app(settings)) as client:
        repaired = _upload(client, content)

    assert repaired.status_code == 201
    assert repaired.json()["data"]["deduplicated"] is True
    assert repaired.json()["data"]["chunk_count"] == 2
    assert repaired.json()["data"]["attempt_count"] == 2


def test_child_hits_expand_to_deduplicated_parent_contexts(tmp_path: Path) -> None:
    long_body = (
        "支付成功后会员会立即生效，请先刷新页面。"
        "如果仍未到账，请核对订单号和当前账号。"
    ) * 5
    with TestClient(create_app(_settings(tmp_path))) as client:
        uploaded = _upload(
            client,
            f"# 大会员\n\n{long_body}\n\n退款规则以结算页面为准。".encode(),
        )
        version_id = uploaded.json()["data"]["version"]["id"]
        chunks = client.get(
            f"/api/v1/knowledge/versions/{version_id}/chunks",
            headers=_headers(),
        ).json()["data"]
        children = [chunk for chunk in chunks if chunk["kind"] == "child"]
        first_parent_children = [
            chunk
            for chunk in children
            if chunk["parent_chunk_id"] == children[0]["parent_chunk_id"]
        ]
        other_parent_child = next(
            chunk
            for chunk in children
            if chunk["parent_chunk_id"] != children[0]["parent_chunk_id"]
        )
        hits = [
            {"chunk_id": first_parent_children[0]["id"], "score": 0.82},
            {"chunk_id": other_parent_child["id"], "score": 0.79},
            {"chunk_id": first_parent_children[1]["id"], "score": 0.91},
        ]
        expanded = client.post(
            f"/api/v1/knowledge/versions/{version_id}/chunks/expand",
            headers=_headers(),
            json={"hits": hits},
        )

    assert expanded.status_code == 200
    contexts = expanded.json()["data"]
    assert len(contexts) == 2
    assert contexts[0]["parent"]["id"] == children[0]["parent_chunk_id"]
    assert contexts[0]["matched_child_ids"] == [
        first_parent_children[0]["id"],
        first_parent_children[1]["id"],
    ]
    assert contexts[0]["best_child_score"] == 0.91
    assert contexts[0]["first_child_rank"] == 1
    assert contexts[1]["parent"]["id"] == other_parent_child["parent_chunk_id"]
    assert contexts[1]["first_child_rank"] == 2


def test_small_to_big_rejects_parent_or_foreign_child_id(tmp_path: Path) -> None:
    with TestClient(create_app(_settings(tmp_path))) as client:
        uploaded = _upload(client, b"# Membership\n\nMember benefits.")
        version_id = uploaded.json()["data"]["version"]["id"]
        parent = client.get(
            f"/api/v1/knowledge/versions/{version_id}/chunks",
            headers=_headers(),
            params={"kind": "parent"},
        ).json()["data"][0]
        response = client.post(
            f"/api/v1/knowledge/versions/{version_id}/chunks/expand",
            headers=_headers(),
            json={
                "hits": [
                    {"chunk_id": parent["id"], "score": 0.8},
                    {"chunk_id": "not-in-this-version", "score": 0.7},
                ]
            },
        )

    assert response.status_code == 409
    assert response.json()["error"]["details"]["invalid_chunk_ids"] == [
        parent["id"],
        "not-in-this-version",
    ]


def test_chunk_debug_runs_specialized_strategy_without_persistence(
    tmp_path: Path,
) -> None:
    settings = _settings(tmp_path)
    with TestClient(create_app(settings)) as client:
        debugged = client.post(
            "/api/v1/knowledge/chunks/debug",
            headers=_headers(),
            json={
                "knowledge_type": "faq",
                "blocks": [
                    {
                        "ordinal": 0,
                        "block_type": "paragraph",
                        "content": "Q：连续包月怎么取消？",
                        "heading_path": ["客服FAQ"],
                    },
                    {
                        "ordinal": 1,
                        "block_type": "paragraph",
                        "content": "A：进入原支付渠道取消。",
                        "heading_path": ["客服FAQ"],
                    },
                    {
                        "ordinal": 2,
                        "block_type": "paragraph",
                        "content": "关键词：自动续费、取消订阅",
                        "heading_path": ["客服FAQ"],
                    },
                ],
            },
        )
        documents = client.get(
            "/api/v1/knowledge/documents",
            headers=_headers(),
        )

    assert debugged.status_code == 200
    result = debugged.json()["data"]
    assert result["parent_count"] == 1
    assert result["child_count"] == 1
    assert result["strategy_counts"] == {"faq": 2}
    assert result["unrepresented_source_ordinals"] == []
    assert result["chunks"][1]["metadata"]["keywords"] == [
        "自动续费",
        "取消订阅",
    ]
    assert documents.json()["data"] == []
