"""PDF、DOCX、Markdown与TXT的结构化Loader。"""

from __future__ import annotations

import re
from io import BytesIO

import pymupdf
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from bili_support.knowledge.loaders.base import DocumentLoadError
from bili_support.knowledge.table_normalization import normalize_table
from bili_support.knowledge.types import (
    LoadedDocument,
    LoadedSourceBlock,
    SourceBlockType,
)

_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MARKDOWN_TABLE_SEPARATOR = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


class PdfLoader:
    """按页提取 PDF 文本块和表格，页码是后续引用定位的核心元数据。"""

    extensions = frozenset({".pdf"})

    def load(
            self,
            *,
            content: bytes,
            filename: str,
            media_type: str,
    ) -> LoadedDocument:
        if not content.startswith(b"%PDF"):
            raise DocumentLoadError("DOCUMENT_SIGNATURE_MISMATCH")
        document: pymupdf.Document = pymupdf.open(  # type: ignore[no-untyped-call]
            stream=content,
            filetype="pdf",
        )
        if document.needs_pass:
            document.close()  # type: ignore[no-untyped-call]
            raise DocumentLoadError("DOCUMENT_ENCRYPTED")
        blocks: list[LoadedSourceBlock] = []
        page_count = document.page_count
        for page_offset in range(document.page_count):
            # PyMuPDF 页下标从 0 开始，对外引用页码从 1 开始。
            page_index = page_offset + 1
            page: pymupdf.Page = document.load_page(  # type: ignore[no-untyped-call]
                page_offset
            )
            for raw_block in page.get_text(  # type: ignore[no-untyped-call]
                    "blocks",
                    sort=True,
            ):
                text = str(raw_block[4]).strip()
                if not text:
                    continue
                blocks.append(
                    LoadedSourceBlock(
                        ordinal=len(blocks),
                        block_type=SourceBlockType.PARAGRAPH,
                        content=text,
                        page_number=page_index,
                    )
                )
            try:
                # 表格单独保存为结构块，不能与页面纯文本混在一起后失去行列关系。
                tables = page.find_tables().tables  # type: ignore[no-untyped-call]
            except (AttributeError, RuntimeError):
                tables = ()
            for table_index, table in enumerate(tables):
                normalized = normalize_table(table.extract())
                if normalized:
                    blocks.append(
                        LoadedSourceBlock(
                            ordinal=len(blocks),
                            block_type=SourceBlockType.TABLE,
                            content=normalized,
                            page_number=page_index,
                            metadata={"table_index": table_index},
                        )
                    )
        result = LoadedDocument(
            filename=filename,
            media_type=media_type,
            blocks=tuple(blocks),
            metadata={"page_count": page_count},
        )
        document.close()  # type: ignore[no-untyped-call]
        return result


class DocxLoader:
    """按 Word 正文顺序提取标题、段落、列表和表格。"""

    extensions = frozenset({".docx"})

    def load(
            self,
            *,
            content: bytes,
            filename: str,
            media_type: str,
    ) -> LoadedDocument:
        if not content.startswith(b"PK"):
            raise DocumentLoadError("DOCUMENT_SIGNATURE_MISMATCH")
        document = DocxDocument(BytesIO(content))
        blocks: list[LoadedSourceBlock] = []
        heading_path: list[str] = []
        for item in document.iter_inner_content():
            # iter_inner_content 同时遍历段落和表格，避免分别读取后破坏原文顺序。
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                style_name = item.style.name.casefold() if item.style else ""
                heading_level = _heading_level(style_name)
                block_type = SourceBlockType.PARAGRAPH
                if heading_level is not None:
                    # 遇到同级或更高级标题时截断旧路径，再加入当前标题。
                    heading_path = heading_path[: heading_level - 1]
                    heading_path.append(text)
                    block_type = SourceBlockType.HEADING
                elif style_name.startswith("list"):
                    block_type = SourceBlockType.LIST
                blocks.append(
                    LoadedSourceBlock(
                        ordinal=len(blocks),
                        block_type=block_type,
                        content=text,
                        heading_path=tuple(heading_path),
                    )
                )
            elif isinstance(item, Table):
                normalized = normalize_table(
                    [[cell.text for cell in row.cells] for row in item.rows]
                )
                if normalized:
                    blocks.append(
                        LoadedSourceBlock(
                            ordinal=len(blocks),
                            block_type=SourceBlockType.TABLE,
                            content=normalized,
                            heading_path=tuple(heading_path),
                        )
                    )
        return LoadedDocument(
            filename=filename,
            media_type=media_type,
            blocks=tuple(blocks),
        )


class MarkdownLoader:
    """解析知识库常用 Markdown 结构，不在 Loader 阶段执行检索分块。"""

    extensions = frozenset({".md", ".markdown"})

    def load(
            self,
            *,
            content: bytes,
            filename: str,
            media_type: str,
    ) -> LoadedDocument:
        text = _decode_text(content)
        lines = text.splitlines()
        blocks: list[LoadedSourceBlock] = []
        heading_path: list[str] = []
        paragraph: list[str] = []
        index = 0

        def flush_paragraph() -> None:
            # 空行、标题和表格都会触发 flush，保证块的原始顺序稳定。
            if not paragraph:
                return
            value = "\n".join(paragraph).strip()
            paragraph.clear()
            if value:
                blocks.append(
                    LoadedSourceBlock(
                        ordinal=len(blocks),
                        block_type=SourceBlockType.PARAGRAPH,
                        content=value,
                        heading_path=tuple(heading_path),
                    )
                )

        while index < len(lines):
            line = lines[index]
            heading = _MARKDOWN_HEADING.match(line)
            if heading:
                flush_paragraph()
                level = len(heading.group(1))
                title = heading.group(2).strip()
                heading_path = heading_path[: level - 1]
                heading_path.append(title)
                blocks.append(
                    LoadedSourceBlock(
                        ordinal=len(blocks),
                        block_type=SourceBlockType.HEADING,
                        content=title,
                        heading_path=tuple(heading_path),
                    )
                )
                index += 1
                continue
            if (
                    "|" in line
                    and index + 1 < len(lines)
                    and _MARKDOWN_TABLE_SEPARATOR.match(lines[index + 1])
            ):
                # Markdown 表格的第二行必须是 --- 分隔符，避免把普通竖线文本误判为表格。
                flush_paragraph()
                table_lines = [line]
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    table_lines.append(lines[index])
                    index += 1
                rows = [_markdown_row(value) for value in table_lines]
                normalized = normalize_table(rows)
                if normalized:
                    blocks.append(
                        LoadedSourceBlock(
                            ordinal=len(blocks),
                            block_type=SourceBlockType.TABLE,
                            content=normalized,
                            heading_path=tuple(heading_path),
                        )
                    )
                continue
            if not line.strip():
                flush_paragraph()
            else:
                paragraph.append(line)
            index += 1
        flush_paragraph()
        return LoadedDocument(
            filename=filename,
            media_type=media_type,
            blocks=tuple(blocks),
        )


class TextLoader:
    """兼容 UTF-8/GB18030 的纯文本 Loader，以空行作为自然段边界。"""

    extensions = frozenset({".txt"})

    def load(
            self,
            *,
            content: bytes,
            filename: str,
            media_type: str,
    ) -> LoadedDocument:
        text = _decode_text(content)
        paragraphs = tuple(
            value.strip()
            for value in re.split(r"\n\s*\n", text)
            if value.strip()
        )
        return LoadedDocument(
            filename=filename,
            media_type=media_type,
            blocks=tuple(
                LoadedSourceBlock(
                    ordinal=index,
                    block_type=SourceBlockType.PARAGRAPH,
                    content=value,
                )
                for index, value in enumerate(paragraphs)
            ),
        )


def _decode_text(content: bytes) -> str:
    # 中文企业存量文档中常见 GB18030，因此在 UTF-8 失败后提供明确回退。
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError("DOCUMENT_ENCODING_UNSUPPORTED")


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"(?:heading|标题)\s*(\d+)", style_name)
    if match is None:
        return None
    return max(1, min(6, int(match.group(1))))


def _markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]
